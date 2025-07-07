"""
Generador de Informe Académico - Implementación de CI/CD con Flask
================================================================

Este script genera un informe académico formal en formato Word (.docx)
documentando todo el proceso de desarrollo e implementación de CI/CD.

Autor: Sistema de Desarrollo
Fecha: 6 de Julio, 2025
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import datetime
import os

def crear_informe_cicd():
    """Genera el informe académico completo en formato Word"""
    
    # Crear documento
    doc = Document()
    
    # ==========================================
    # PORTADA
    # ==========================================
    
    # Título principal
    titulo = doc.add_heading('INFORME TÉCNICO', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitulo = doc.add_heading('Implementación de CI/CD para Aplicación Web Flask\ncon GitHub Actions y Render', level=1)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del proyecto
    doc.add_paragraph()
    doc.add_paragraph()
    
    info_proyecto = doc.add_paragraph()
    info_proyecto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_proyecto.add_run('Proyecto: ').bold = True
    info_proyecto.add_run('Flask Todo Application\n')
    info_proyecto.add_run('Autor: ').bold = True
    info_proyecto.add_run('joe-696\n')
    info_proyecto.add_run('Fecha: ').bold = True
    info_proyecto.add_run(f'{datetime.datetime.now().strftime("%d de %B, %Y")}\n')
    info_proyecto.add_run('Repositorio: ').bold = True
    info_proyecto.add_run('https://github.com/joe-696/flask-todo-app\n')
    info_proyecto.add_run('URL Producción: ').bold = True
    info_proyecto.add_run('https://flask-todo-app-joe.onrender.com')
    
    # Salto de página
    doc.add_page_break()
    
    # ==========================================
    # ÍNDICE
    # ==========================================
    
    doc.add_heading('ÍNDICE', level=1)
    
    indice_items = [
        "1. RESUMEN EJECUTIVO",
        "2. INTRODUCCIÓN", 
        "3. OBJETIVOS",
        "4. MARCO TEÓRICO",
        "   4.1. CI/CD (Continuous Integration/Continuous Deployment)",
        "   4.2. GitHub Actions",
        "   4.3. Flask Framework",
        "   4.4. Testing Automatizado",
        "5. METODOLOGÍA",
        "6. DESARROLLO DEL PROYECTO",
        "   6.1. Estructura de la Aplicación Flask",
        "   6.2. Configuración del Repositorio Git",
        "   6.3. Implementación de Tests Automáticos",
        "   6.4. Configuración de GitHub Actions",
        "   6.5. Despliegue en Render",
        "7. ARQUITECTURA DEL SISTEMA",
        "8. FLUJO CI/CD IMPLEMENTADO",
        "9. RESULTADOS OBTENIDOS",
        "10. CAPTURAS DE PANTALLA",
        "11. CONCLUSIONES",
        "12. RECOMENDACIONES",
        "13. BIBLIOGRAFÍA",
        "14. ANEXOS"
    ]
    
    for item in indice_items:
        p = doc.add_paragraph(item)
        if item.startswith('   '):
            p.style = 'List Number'
        else:
            p.style = 'List Number'
    
    doc.add_page_break()
    
    # ==========================================
    # 1. RESUMEN EJECUTIVO
    # ==========================================
    
    doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
    
    doc.add_paragraph(
        "Este informe documenta la implementación exitosa de un pipeline de CI/CD "
        "(Continuous Integration/Continuous Deployment) para una aplicación web "
        "desarrollada con Flask. El proyecto abarca desde la creación de una "
        "aplicación CRUD simple hasta la implementación de un sistema automatizado "
        "de integración y despliegue continuo utilizando GitHub Actions y Render."
    )
    
    doc.add_paragraph(
        "Los resultados obtenidos incluyen: una aplicación web funcional con "
        "operaciones CRUD completas, tests automáticos que validan la funcionalidad, "
        "un pipeline de CI/CD que ejecuta automáticamente en cada cambio de código, "
        "y despliegue automático a producción con validación previa."
    )
    
    # ==========================================
    # 2. INTRODUCCIÓN
    # ==========================================
    
    doc.add_heading('2. INTRODUCCIÓN', level=1)
    
    doc.add_paragraph(
        "En el desarrollo de software moderno, la implementación de prácticas "
        "DevOps como CI/CD es fundamental para mantener la calidad del código "
        "y acelerar los ciclos de desarrollo. Este proyecto demuestra la "
        "implementación práctica de estas metodologías en una aplicación real."
    )
    
    doc.add_paragraph(
        "La aplicación desarrollada es un sistema de gestión de tareas (Todo List) "
        "que permite a los usuarios crear, leer, actualizar y eliminar tareas. "
        "Aunque simple en concepto, sirve como base perfecta para demostrar "
        "conceptos avanzados de CI/CD."
    )
    
    # ==========================================
    # 3. OBJETIVOS
    # ==========================================
    
    doc.add_heading('3. OBJETIVOS', level=1)
    
    doc.add_heading('3.1. Objetivo General', level=2)
    doc.add_paragraph(
        "Implementar un sistema completo de CI/CD para una aplicación web Flask, "
        "incluyendo tests automáticos, validación de código y despliegue automático "
        "a producción."
    )
    
    doc.add_heading('3.2. Objetivos Específicos', level=2)
    objetivos = [
        "Desarrollar una aplicación web CRUD funcional utilizando Flask",
        "Implementar tests automáticos con pytest",
        "Configurar GitHub Actions para CI/CD",
        "Establecer despliegue automático en Render",
        "Documentar todo el proceso de implementación",
        "Validar el funcionamiento del pipeline completo"
    ]
    
    for objetivo in objetivos:
        doc.add_paragraph(f"• {objetivo}")
    
    # ==========================================
    # 4. MARCO TEÓRICO
    # ==========================================
    
    doc.add_heading('4. MARCO TEÓRICO', level=1)
    
    doc.add_heading('4.1. CI/CD (Continuous Integration/Continuous Deployment)', level=2)
    doc.add_paragraph(
        "CI/CD es una metodología de desarrollo que permite a los equipos "
        "entregar cambios de código de forma rápida y confiable. La Integración "
        "Continua (CI) implica la fusión frecuente de cambios de código en un "
        "repositorio central, mientras que el Despliegue Continuo (CD) automatiza "
        "la entrega de aplicaciones a entornos de producción."
    )
    
    doc.add_heading('4.2. GitHub Actions', level=2)
    doc.add_paragraph(
        "GitHub Actions es una plataforma de CI/CD que permite automatizar "
        "workflows directamente desde el repositorio de GitHub. Utiliza archivos "
        "YAML para definir pipelines que se ejecutan en respuesta a eventos "
        "como commits, pull requests, o schedules."
    )
    
    doc.add_heading('4.3. Flask Framework', level=2)
    doc.add_paragraph(
        "Flask es un micro-framework web para Python que proporciona las "
        "herramientas básicas para construir aplicaciones web. Es minimalista "
        "pero extensible, lo que lo hace ideal para proyectos que requieren "
        "flexibilidad y control granular."
    )
    
    doc.add_heading('4.4. Testing Automatizado', level=2)
    doc.add_paragraph(
        "Los tests automáticos son scripts que verifican el comportamiento "
        "esperado del código. En Python, pytest es una de las herramientas "
        "más populares para escribir y ejecutar tests, proporcionando "
        "funcionalidades avanzadas como fixtures y parametrización."
    )
    
    # ==========================================
    # 5. METODOLOGÍA
    # ==========================================
    
    doc.add_heading('5. METODOLOGÍA', level=1)
    
    doc.add_paragraph(
        "El desarrollo del proyecto siguió una metodología incremental, "
        "implementando funcionalidades en el siguiente orden:"
    )
    
    metodologia = [
        "Análisis de requerimientos y diseño de la aplicación",
        "Desarrollo de la aplicación Flask básica",
        "Implementación de operaciones CRUD",
        "Creación de templates HTML con Bootstrap",
        "Configuración del entorno de desarrollo",
        "Implementación de tests automáticos",
        "Configuración de GitHub Actions",
        "Despliegue en plataforma Render",
        "Validación del pipeline CI/CD completo",
        "Documentación y análisis de resultados"
    ]
    
    for i, paso in enumerate(metodologia, 1):
        doc.add_paragraph(f"{i}. {paso}")
    
    # ==========================================
    # 6. DESARROLLO DEL PROYECTO
    # ==========================================
    
    doc.add_heading('6. DESARROLLO DEL PROYECTO', level=1)
    
    doc.add_heading('6.1. Estructura de la Aplicación Flask', level=2)
    doc.add_paragraph(
        "La aplicación fue estructurada siguiendo las mejores prácticas de Flask:"
    )
    
    estructura_codigo = '''
proyecto/
├── app.py                  # Aplicación principal
├── requirements.txt        # Dependencias
├── Procfile               # Configuración Render
├── .gitignore            # Archivos ignorados por Git
├── templates/            # Plantillas HTML
│   ├── base.html         # Template base
│   ├── index.html        # Lista de tareas
│   ├── add_todo.html     # Agregar tarea
│   └── edit_todo.html    # Editar tarea
├── tests/                # Tests automáticos
│   ├── __init__.py       
│   └── test_app.py       # Tests principales
└── .github/              # GitHub Actions
    └── workflows/
        └── ci-cd.yml     # Pipeline CI/CD
'''
    
    p = doc.add_paragraph()
    p.style = 'Intense Quote'
    p.add_run(estructura_codigo)
    
    doc.add_heading('6.2. Configuración del Repositorio Git', level=2)
    doc.add_paragraph(
        "Se configuró un repositorio Git con las siguientes características:"
    )
    
    git_config = [
        "Inicialización del repositorio local",
        "Configuración de credenciales de usuario",
        "Creación de .gitignore apropiado",
        "Conexión con repositorio remoto en GitHub",
        "Configuración de token de acceso personal con permisos de workflow"
    ]
    
    for config in git_config:
        doc.add_paragraph(f"• {config}")
    
    doc.add_heading('6.3. Implementación de Tests Automáticos', level=2)
    doc.add_paragraph(
        "Se implementaron tests comprehensivos que cubren:"
    )
    
    tests_implementados = [
        "Test de carga de página principal",
        "Test de funcionalidad de creación de tareas",
        "Test de validación de campos obligatorios",
        "Test de edición de tareas existentes",
        "Test de eliminación de tareas",
        "Test de cambio de estado (completado/pendiente)",
        "Test del modelo de datos Todo"
    ]
    
    for test in tests_implementados:
        doc.add_paragraph(f"• {test}")
    
    doc.add_heading('6.4. Configuración de GitHub Actions', level=2)
    doc.add_paragraph(
        "El workflow de GitHub Actions incluye dos jobs principales:"
    )
    
    doc.add_paragraph().add_run("Job de Testing:").bold = True
    test_steps = [
        "Checkout del código fuente",
        "Configuración de Python 3.11",
        "Instalación de dependencias",
        "Ejecución de tests con pytest",
        "Validación de startup de la aplicación"
    ]
    
    for step in test_steps:
        doc.add_paragraph(f"  • {step}")
    
    doc.add_paragraph().add_run("Job de Deploy:").bold = True
    deploy_steps = [
        "Ejecución condicional (solo en branch main)",
        "Dependencia del job de testing",
        "Notificación de despliegue automático por Render"
    ]
    
    for step in deploy_steps:
        doc.add_paragraph(f"  • {step}")
    
    doc.add_heading('6.5. Despliegue en Render', level=2)
    doc.add_paragraph(
        "La configuración de Render incluye:"
    )
    
    render_config = [
        "Conexión automática con repositorio GitHub",
        "Detección automática de aplicación Flask",
        "Build command: pip install -r requirements.txt",
        "Start command: gunicorn app:app",
        "Variables de entorno para producción",
        "Despliegue automático en cada push a main"
    ]
    
    for config in render_config:
        doc.add_paragraph(f"• {config}")
    
    # ==========================================
    # 7. ARQUITECTURA DEL SISTEMA
    # ==========================================
    
    doc.add_heading('7. ARQUITECTURA DEL SISTEMA', level=1)
    
    doc.add_paragraph(
        "La arquitectura implementada sigue un patrón de tres capas:"
    )
    
    doc.add_paragraph().add_run("Capa de Presentación:").bold = True
    doc.add_paragraph(
        "Templates HTML con Bootstrap para interfaz responsive y moderna."
    )
    
    doc.add_paragraph().add_run("Capa de Lógica de Negocio:").bold = True
    doc.add_paragraph(
        "Aplicación Flask con rutas para operaciones CRUD y manejo de formularios."
    )
    
    doc.add_paragraph().add_run("Capa de Datos:").bold = True
    doc.add_paragraph(
        "SQLAlchemy ORM con base de datos SQLite para desarrollo y PostgreSQL "
        "disponible para producción."
    )
    
    doc.add_paragraph(
        "[INSERTAR DIAGRAMA DE ARQUITECTURA]"
    )
    
    # ==========================================
    # 8. FLUJO CI/CD IMPLEMENTADO
    # ==========================================
    
    doc.add_heading('8. FLUJO CI/CD IMPLEMENTADO', level=1)
    
    doc.add_paragraph(
        "El flujo de CI/CD implementado sigue el siguiente proceso:"
    )
    
    flujo_cicd = [
        "1. Developer hace git push al repositorio",
        "2. GitHub detecta cambios y activa workflow",
        "3. GitHub Actions inicia job de testing",
        "4. Se ejecutan todos los tests automáticos",
        "5. Si tests pasan: continúa al deploy",
        "6. Si tests fallan: detiene el proceso",
        "7. Render detecta cambios en GitHub",
        "8. Render ejecuta build automático",
        "9. Render despliega nueva versión",
        "10. Aplicación actualizada en producción"
    ]
    
    for paso in flujo_cicd:
        doc.add_paragraph(paso)
    
    doc.add_paragraph(
        "[INSERTAR DIAGRAMA DE FLUJO CI/CD]"
    )
    
    # ==========================================
    # 9. RESULTADOS OBTENIDOS
    # ==========================================
    
    doc.add_heading('9. RESULTADOS OBTENIDOS', level=1)
    
    doc.add_paragraph().add_run("Aplicación Funcional:").bold = True
    doc.add_paragraph(
        "✅ Aplicación web completamente funcional\n"
        "✅ Operaciones CRUD implementadas\n"
        "✅ Interfaz responsive con Bootstrap\n"
        "✅ Validación de formularios\n"
        "✅ Mensajes de feedback al usuario"
    )
    
    doc.add_paragraph().add_run("CI/CD Automatizado:").bold = True
    doc.add_paragraph(
        "✅ Pipeline de GitHub Actions funcionando\n"
        "✅ Tests automáticos ejecutándose en cada push\n"
        "✅ Deploy automático a producción\n"
        "✅ Validación previa antes del despliegue\n"
        "✅ Tiempo total de pipeline: ~21 segundos"
    )
    
    doc.add_paragraph().add_run("Calidad del Código:").bold = True
    doc.add_paragraph(
        "✅ Cobertura de tests implementada\n"
        "✅ Validación automática de funcionalidades\n"
        "✅ Detección temprana de errores\n"
        "✅ Código versionado y documentado"
    )
    
    # ==========================================
    # 10. CAPTURAS DE PANTALLA
    # ==========================================
    
    doc.add_heading('10. CAPTURAS DE PANTALLA', level=1)
    
    capturas = [
        "Aplicación en funcionamiento - Página principal",
        "GitHub Actions - Workflow ejecutándose",
        "GitHub Actions - Tests exitosos",
        "Render - Dashboard de despliegue",
        "Aplicación en producción - URL live"
    ]
    
    for i, captura in enumerate(capturas, 1):
        doc.add_paragraph(f"Figura {i}: {captura}")
        doc.add_paragraph("[INSERTAR CAPTURA AQUÍ]")
        doc.add_paragraph()
    
    # ==========================================
    # 11. CONCLUSIONES
    # ==========================================
    
    doc.add_heading('11. CONCLUSIONES', level=1)
    
    conclusiones = [
        "La implementación de CI/CD mejora significativamente la calidad y "
        "confiabilidad del software desarrollado.",
        
        "GitHub Actions proporciona una plataforma robusta y accesible para "
        "implementar pipelines de CI/CD sin costo adicional.",
        
        "Los tests automáticos son fundamentales para detectar errores antes "
        "de que lleguen a producción.",
        
        "El despliegue automático reduce errores humanos y acelera el tiempo "
        "de entrega de nuevas funcionalidades.",
        
        "La combinación Flask + GitHub Actions + Render proporciona un stack "
        "completo para desarrollo web moderno.",
        
        "La documentación y versionado del código facilita el mantenimiento "
        "y colaboración en equipo."
    ]
    
    for conclusion in conclusiones:
        doc.add_paragraph(f"• {conclusion}")
    
    # ==========================================
    # 12. RECOMENDACIONES
    # ==========================================
    
    doc.add_heading('12. RECOMENDACIONES', level=1)
    
    recomendaciones = [
        "Implementar métricas de cobertura de código para mejorar la calidad de tests",
        "Agregar múltiples entornos (desarrollo, staging, producción)",
        "Implementar notificaciones automáticas en caso de fallos",
        "Considerar implementar análisis estático de código (linting)",
        "Evaluar la migración a PostgreSQL para mayor robustez en producción",
        "Implementar monitoring y logging avanzado",
        "Considerar la implementación de blue-green deployment"
    ]
    
    for recomendacion in recomendaciones:
        doc.add_paragraph(f"• {recomendacion}")
    
    # ==========================================
    # 13. BIBLIOGRAFÍA
    # ==========================================
    
    doc.add_heading('13. BIBLIOGRAFÍA', level=1)
    
    bibliografia = [
        "Flask Documentation. (2025). Flask Web Development Framework. https://flask.palletsprojects.com/",
        "GitHub Actions Documentation. (2025). Automating workflows with GitHub Actions. https://docs.github.com/actions",
        "Render Documentation. (2025). Deploy and scale applications. https://render.com/docs",
        "pytest Documentation. (2025). Testing framework for Python. https://docs.pytest.org/",
        "SQLAlchemy Documentation. (2025). Python SQL toolkit. https://www.sqlalchemy.org/",
        "Bootstrap Documentation. (2025). Frontend framework. https://getbootstrap.com/"
    ]
    
    for ref in bibliografia:
        doc.add_paragraph(ref)
    
    # ==========================================
    # 14. ANEXOS
    # ==========================================
    
    doc.add_heading('14. ANEXOS', level=1)
    
    doc.add_heading('Anexo A: Código Fuente Principal (app.py)', level=2)
    doc.add_paragraph("[Incluir código fuente completo de app.py]")
    
    doc.add_heading('Anexo B: Archivo de Workflow GitHub Actions', level=2)
    doc.add_paragraph("[Incluir contenido completo de ci-cd.yml]")
    
    doc.add_heading('Anexo C: Tests Implementados', level=2)
    doc.add_paragraph("[Incluir código completo de test_app.py]")
    
    doc.add_heading('Anexo D: Configuración de Despliegue', level=2)
    doc.add_paragraph("[Incluir Procfile y requirements.txt]")
    
    # Guardar documento
    filename = f"Informe_CICD_Flask_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    
    return filename

if __name__ == "__main__":
    print("🔄 Generando informe académico...")
    
    try:
        archivo = crear_informe_cicd()
        print(f"✅ ¡Informe generado exitosamente!")
        print(f"📄 Archivo: {archivo}")
        print(f"📁 Ubicación: {os.path.abspath(archivo)}")
        print("\n📋 El informe incluye:")
        print("   • Portada e índice profesional")
        print("   • Marco teórico completo")
        print("   • Metodología detallada")
        print("   • Resultados y análisis")
        print("   • Secciones para capturas de pantalla")
        print("   • Conclusiones y recomendaciones")
        print("   • Bibliografía y anexos")
        print("\n💡 Nota: Recuerda agregar las capturas de pantalla en las secciones marcadas")
        
    except ImportError as e:
        print("❌ Error: Falta instalar la librería python-docx")
        print("💻 Ejecuta: pip install python-docx")
        print(f"   Error técnico: {e}")
    except Exception as e:
        print(f"❌ Error inesperado: {e}")
        print("🔧 Verifica que tengas permisos de escritura en el directorio")
